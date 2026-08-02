#!/usr/bin/env python3
"""
换墨 — 文档格式转换器
支持 Markdown / TXT / DOCX → PDF，以及 PDF 合并
启动后自动打开浏览器，拖拽文件即可转换。
"""

from __future__ import annotations

import base64
import logging
import os
import re
import sys
import tempfile
import threading
import time
import webbrowser
from io import BytesIO
from pathlib import Path

# ── Dependencies ──────────────────────────────────────────────
try:
    from flask import Flask, request, jsonify, send_file, render_template_string
except ImportError:
    sys.exit("需要 Flask: pip install flask")

try:
    from markdown_it import MarkdownIt
except ImportError:
    sys.exit("需要 markdown-it-py: pip install markdown-it-py")

try:
    from fpdf import FPDF, TextStyle
except ImportError:
    sys.exit("需要 fpdf2: pip install fpdf2")

try:
    from pypdf import PdfReader, PdfWriter
except ImportError:
    try:
        from PyPDF2 import PdfReader, PdfWriter
    except ImportError:
        PdfReader = None
        PdfWriter = None

try:
    import pypdfium2 as pdfium
except ImportError:
    pdfium = None

# ── Logging ───────────────────────────────────────────────────
logging.getLogger("fpdf").setLevel(logging.ERROR)
logging.getLogger("fontTools").setLevel(logging.ERROR)
logging.getLogger("fontTools.subset").setLevel(logging.ERROR)
logging.getLogger("werkzeug").setLevel(logging.WARNING)

log = logging.getLogger("md2pdf-web")

# ── Cross-platform CJK font detection ─────────────────────────
FONT_CANDIDATES: dict[str, list[str]] = {
    "darwin": [
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/PingFang.ttc",
    ],
    "linux": [
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
    ],
    "win32": [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/mingliu.ttc",
    ],
}

_FPDF_FONT_NAME = "CJK"
_md_parser = MarkdownIt()


def detect_cjk_font() -> str | None:
    for path in FONT_CANDIDATES.get(sys.platform, []):
        if os.path.isfile(path):
            return path
    return None


def md_to_clean_html(text: str) -> str:
    html = _md_parser.render(text)
    html = re.sub(r"</?thead>", "", html)
    html = re.sub(r"</?tbody>", "", html)
    html = re.sub(r"<pre><code>", "<pre>", html)
    html = re.sub(r"</code></pre>", "</pre>", html)
    html = html.replace("<table>", '<table border="1">')
    return html


def convert_md_to_pdf_bytes(md_text: str, title: str, font_path: str) -> bytes:
    """Convert markdown text to PDF, return bytes."""
    html_body = md_to_clean_html(md_text)
    full_html = f"<h1>{title}</h1><hr>{html_body}"

    pdf = FPDF()
    pdf.add_font(_FPDF_FONT_NAME, style="", fname=font_path)
    pdf.add_font(_FPDF_FONT_NAME, style="B", fname=font_path)
    pdf.add_font(_FPDF_FONT_NAME, style="I", fname=font_path)
    pdf.add_font(_FPDF_FONT_NAME, style="BI", fname=font_path)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_left_margin(25)
    pdf.set_right_margin(25)
    pdf.add_page()

    # Suppress fonttools noise
    import io as _io
    _stderr_ctx = (
        __import__("contextlib").redirect_stderr(_io.StringIO())
        if not os.environ.get("MD2PDF_DEBUG")
        else __import__("contextlib").nullcontext()
    )
    with _stderr_ctx:
        pdf.write_html(
            full_html,
            font_family=_FPDF_FONT_NAME,
            tag_styles={
                "pre": TextStyle(font_family=_FPDF_FONT_NAME, font_size_pt=10),
                "code": TextStyle(font_family=_FPDF_FONT_NAME, font_size_pt=10),
            },
        )

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    try:
        pdf.output(tmp.name)
        tmp.close()
        return Path(tmp.name).read_bytes()
    finally:
        os.unlink(tmp.name)


def convert_txt_to_pdf_bytes(text: str, title: str, font_path: str) -> bytes:
    """Convert plain text to PDF, preserving whitespace."""
    escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    full_html = f"<h1>{title}</h1><hr><pre>{escaped}</pre>"

    pdf = FPDF()
    pdf.add_font(_FPDF_FONT_NAME, style="", fname=font_path)
    pdf.add_font(_FPDF_FONT_NAME, style="B", fname=font_path)
    pdf.add_font(_FPDF_FONT_NAME, style="I", fname=font_path)
    pdf.add_font(_FPDF_FONT_NAME, style="BI", fname=font_path)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_left_margin(25)
    pdf.set_right_margin(25)
    pdf.add_page()

    import io as _io
    _stderr_ctx = (
        __import__("contextlib").redirect_stderr(_io.StringIO())
        if not os.environ.get("MD2PDF_DEBUG")
        else __import__("contextlib").nullcontext()
    )
    with _stderr_ctx:
        pdf.write_html(
            full_html,
            font_family=_FPDF_FONT_NAME,
            tag_styles={
                "pre": TextStyle(font_family=_FPDF_FONT_NAME, font_size_pt=11),
            },
        )

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    try:
        pdf.output(tmp.name)
        tmp.close()
        return Path(tmp.name).read_bytes()
    finally:
        os.unlink(tmp.name)


def convert_docx_to_pdf_bytes(docx_bytes: bytes, title: str, font_path: str) -> bytes:
    """Convert a .docx file to PDF via python-docx → HTML."""
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(docx_bytes))

    html_parts = [f"<h1>{title}</h1><hr>"]

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            html_parts.append("<br>")
            continue

        # Detect heading style
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            level = style_name.replace("Heading", "").strip()
            try:
                lv = int(level)
            except ValueError:
                lv = 2
            lv = max(1, min(6, lv))
            html_parts.append(f"<h{lv}>{_escape_html(text)}</h{lv}>")
            continue

        # Inline formatting
        runs_html = []
        for run in para.runs:
            t = _escape_html(run.text)
            if not t:
                continue
            if run.bold:
                t = f"<b>{t}</b>"
            if run.italic:
                t = f"<i>{t}</i>"
            runs_html.append(t)
        para_html = "".join(runs_html) if runs_html else _escape_html(text)
        html_parts.append(f"<p>{para_html}</p>")

    # Tables
    for table in doc.tables:
        html_parts.append('<table border="1">')
        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row.cells:
                html_parts.append(f"<td>{_escape_html(cell.text)}</td>")
            html_parts.append("</tr>")
        html_parts.append("</table>")
        html_parts.append("<br>")

    full_html = "".join(html_parts)

    pdf = FPDF()
    pdf.add_font(_FPDF_FONT_NAME, style="", fname=font_path)
    pdf.add_font(_FPDF_FONT_NAME, style="B", fname=font_path)
    pdf.add_font(_FPDF_FONT_NAME, style="I", fname=font_path)
    pdf.add_font(_FPDF_FONT_NAME, style="BI", fname=font_path)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_left_margin(25)
    pdf.set_right_margin(25)
    pdf.add_page()

    import io as _io
    _stderr_ctx = (
        __import__("contextlib").redirect_stderr(_io.StringIO())
        if not os.environ.get("MD2PDF_DEBUG")
        else __import__("contextlib").nullcontext()
    )
    with _stderr_ctx:
        pdf.write_html(
            full_html,
            font_family=_FPDF_FONT_NAME,
            tag_styles={
                "pre": TextStyle(font_family=_FPDF_FONT_NAME, font_size_pt=11),
            },
        )

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    try:
        pdf.output(tmp.name)
        tmp.close()
        return Path(tmp.name).read_bytes()
    finally:
        os.unlink(tmp.name)


def _escape_html(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def merge_pdf_bytes(pdf_items: list[tuple[str, bytes]]) -> bytes:
    """Merge multiple PDF files into one, preserving page order."""
    if PdfReader is None or PdfWriter is None:
        raise RuntimeError("需要 pypdf: pip install pypdf 或 PyPDF2")

    writer = PdfWriter()
    for name, raw in pdf_items:
        reader = PdfReader(BytesIO(raw))
        if getattr(reader, "is_encrypted", False):
            raise ValueError(f"{name} 已加密，无法合并")
        if not reader.pages:
            raise ValueError(f"{name} 没有可合并的页面")
        for page in reader.pages:
            writer.add_page(page)

    out = BytesIO()
    writer.write(out)
    return out.getvalue()


def merge_selected_pages(pdf_items: list[tuple[str, bytes]], order: list[tuple[str, int]]) -> bytes:
    """Merge only the selected pages, in the given order."""
    if PdfReader is None or PdfWriter is None:
        raise RuntimeError("需要 pypdf: pip install pypdf 或 PyPDF2")

    readers = {name: PdfReader(BytesIO(raw)) for name, raw in pdf_items}
    writer = PdfWriter()
    for name, page_index in order:
        reader = readers[name]
        if getattr(reader, "is_encrypted", False):
            raise ValueError(f"{name} 已加密，无法合并")
        writer.add_page(reader.pages[page_index])

    out = BytesIO()
    writer.write(out)
    return out.getvalue()


def _bitmap_to_png(bitmap) -> bytes:
    """Encode a pypdfium2 bitmap as PNG without requiring Pillow."""
    import struct
    import zlib

    width = bitmap.width
    height = bitmap.height
    stride = bitmap.stride
    channels = bitmap.n_channels
    mode = getattr(bitmap, "mode", "") or ""
    data = bytes(bitmap.buffer)

    swap_rb = mode.startswith("BGR")
    has_alpha = "A" in mode or channels == 4

    raw = bytearray()
    for y in range(height):
        row_start = y * stride
        row = data[row_start:row_start + width * channels]
        raw.append(0)  # PNG filter type: None
        if channels == 4:
            if swap_rb:
                for x in range(width):
                    i = x * 4
                    raw += bytes((row[i + 2], row[i + 1], row[i], row[i + 3]))
            else:
                raw += row
        elif channels == 3:
            if swap_rb:
                for x in range(width):
                    i = x * 3
                    raw += bytes((row[i + 2], row[i + 1], row[i]))
            else:
                raw += row
        else:
            raw += row

    def chunk(tag: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload))
            + tag
            + payload
            + struct.pack(">I", zlib.crc32(tag + payload) & 0xFFFFFFFF)
        )

    color_type = 6 if has_alpha else (2 if channels >= 3 else 0)
    ihdr = struct.pack(">IIBBBBB", width, height, 8, color_type, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(bytes(raw), 6))
        + chunk(b"IEND", b"")
    )


def render_page_preview(pdf_bytes: bytes, page_index: int, max_width: int = 560) -> bytes | None:
    """Render one PDF page as a PNG preview."""
    if pdfium is None:
        return None
    pdf = pdfium.PdfDocument(BytesIO(pdf_bytes))
    try:
        page = pdf[page_index]
        try:
            page_width, _ = page.get_size()
        except Exception:
            page_width = 595.0
        scale = 2.0
        if page_width and page_width * scale > max_width:
            scale = max_width / page_width
        bitmap = page.render(scale=scale)
        return _bitmap_to_png(bitmap)
    finally:
        pdf.close()


# ── Flask App ─────────────────────────────────────────────────
app = Flask(__name__)

# Output directory – persisted in memory during session
OUTPUT_DIR = str(Path.home() / "Downloads")

HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>换墨</title>
<style>
  :root {
    --bg: #f5f5f7;
    --card: #ffffff;
    --text: #1d1d1f;
    --secondary: #86868b;
    --accent: #0071e3;
    --border: #d2d2d7;
    --hover: #f0f0f2;
    --radius: 12px;
  }
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body {
    font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", "PingFang SC", "Helvetica Neue", sans-serif;
    background: var(--bg);
    color: var(--text);
    display: flex; justify-content: center;
    min-height: 100vh; padding: 40px 20px;
    -webkit-font-smoothing: antialiased;
  }
  .container { width: 100%; max-width: 600px; }
  h1 {
    font-size: 28px; font-weight: 700;
    text-align: center; margin-bottom: 4px;
    letter-spacing: -0.5px;
  }
  .subtitle { text-align: center; color: var(--secondary); font-size: 14px; margin-bottom: 32px; }

  /* Mode switch */
  .mode-switch {
    display: flex; gap: 8px; justify-content: center;
    margin-bottom: 20px; background: var(--card);
    border: 1px solid var(--border); border-radius: 10px; padding: 4px;
  }
  .mode-btn {
    flex: 1; padding: 8px 12px; border: none; border-radius: 7px;
    background: transparent; color: var(--secondary);
    font-size: 14px; font-weight: 500; cursor: pointer;
    transition: all 0.15s; white-space: nowrap;
  }
  .mode-btn:hover { color: var(--text); }
  .mode-btn.active { background: var(--accent); color: #fff; }

  /* Drop zone */
  .drop-zone {
    background: var(--card);
    border: 2px dashed var(--border);
    border-radius: var(--radius);
    padding: 48px 24px;
    text-align: center;
    cursor: pointer;
    transition: all 0.2s;
    margin-bottom: 20px;
  }
  .drop-zone:hover, .drop-zone.drag-over {
    border-color: var(--accent);
    background: #f0f4ff;
  }
  .drop-zone .icon { font-size: 40px; margin-bottom: 12px; }
  .drop-zone p { color: var(--secondary); font-size: 15px; }
  .drop-zone .browse { color: var(--accent); font-weight: 500; }

  /* File list */
  .file-list {
    background: var(--card);
    border-radius: var(--radius);
    overflow: hidden;
    margin-bottom: 20px;
    display: none;
  }
  .file-list.has-files { display: block; }
  .file-list-header {
    font-size: 13px; font-weight: 600;
    color: var(--secondary); text-transform: uppercase;
    padding: 14px 20px 8px;
  }
  .file-item {
    display: flex; align-items: center; justify-content: space-between;
    padding: 10px 20px; border-top: 1px solid #f0f0f0;
    font-size: 15px;
  }
  .file-item .name { display: flex; align-items: center; gap: 10px; overflow: hidden; }
  .file-item .name span { white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
  .file-item .remove {
    background: none; border: none; color: #ff3b30;
    cursor: pointer; font-size: 14px; padding: 4px 8px;
    border-radius: 6px; opacity: 0.6; transition: all 0.15s;
  }
  .file-item .remove:hover { opacity: 1; background: #fff0f0; }

  /* Output dir */
  .output-row {
    display: flex; gap: 10px; align-items: center;
    margin-bottom: 20px;
  }
  .output-row label { font-size: 14px; font-weight: 500; white-space: nowrap; }
  .output-row input {
    flex: 1; padding: 10px 14px;
    border: 1px solid var(--border); border-radius: 8px;
    font-size: 14px; background: var(--card);
    font-family: "SF Mono", "Menlo", "Consolas", monospace;
  }
  .output-row input:focus { outline: none; border-color: var(--accent); box-shadow: 0 0 0 3px rgba(0,113,227,0.15); }
  .btn-outline {
    padding: 10px 16px;
    border: 1px solid var(--border); border-radius: 8px;
    background: var(--card); font-size: 14px;
    cursor: pointer; white-space: nowrap;
    transition: all 0.15s;
  }
  .btn-outline:hover { background: var(--hover); }

  /* Convert button */
  .btn-convert {
    width: 100%; padding: 14px;
    background: var(--accent); color: #fff;
    border: none; border-radius: 10px;
    font-size: 17px; font-weight: 600;
    cursor: pointer; transition: all 0.15s;
    letter-spacing: -0.2px;
  }
  .btn-convert:hover { opacity: 0.92; }
  .btn-convert:disabled { opacity: 0.4; cursor: not-allowed; }
  .btn-convert .spinner { display: none; }
  .btn-convert.loading .spinner { display: inline-block; margin-right: 8px; animation: spin 0.8s linear infinite; }
  .btn-convert.loading .btn-text { display: none; }
  @keyframes spin { to { transform: rotate(360deg); } }

  /* Results */
  .results { margin-top: 24px; display: none; }
  .results.has-results { display: block; }
  .results-header {
    font-size: 13px; font-weight: 600;
    color: var(--secondary); text-transform: uppercase;
    margin-bottom: 8px;
  }
  .result-item {
    display: flex; align-items: center; justify-content: space-between;
    background: var(--card); border-radius: 10px;
    padding: 14px 20px; margin-bottom: 8px;
    font-size: 15px;
  }
  .result-item .status { font-size: 18px; margin-right: 10px; }
  .result-item .info { flex: 1; }
  .result-item .info .filename { font-weight: 500; }
  .result-item .info .size { font-size: 13px; color: var(--secondary); }
  .result-item .open-btn {
    padding: 8px 16px; background: var(--accent); color: #fff;
    border: none; border-radius: 6px; font-size: 14px;
    cursor: pointer; text-decoration: none; transition: all 0.15s;
  }
  .result-item .open-btn:hover { opacity: 0.9; }
  .result-item.error .status { color: #ff3b30; }
  .result-item.error .info { color: #ff3b30; }

  /* Page editor (merge mode) */
  .page-editor { display: none; margin-bottom: 20px; }
  .page-editor.visible { display: block; }
  .page-editor-header {
    font-size: 13px; font-weight: 600;
    color: var(--secondary); text-transform: uppercase;
    padding: 14px 20px 8px; background: var(--card);
    border-radius: 10px 10px 0 0; border-bottom: 1px solid #f0f0f0;
  }
  .page-list {
    display: grid; grid-template-columns: repeat(auto-fill, minmax(118px, 1fr));
    gap: 10px; background: var(--card); padding: 14px;
    border-radius: 0 0 10px 10px;
  }
  .page-card {
    border: 1px solid var(--border); border-radius: 8px;
    padding: 6px; cursor: grab; position: relative;
    background: #fff; transition: border-color 0.15s, box-shadow 0.15s;
  }
  .page-card:hover { border-color: var(--accent); }
  .page-card.dragging { opacity: 0.45; }
  .page-card.drop-target { border-color: var(--accent); box-shadow: 0 0 0 2px rgba(0,113,227,0.18); }
  .page-card img, .page-placeholder {
    width: 100%; aspect-ratio: 210 / 297; object-fit: contain;
    background: #fff; border: 1px solid #eee; border-radius: 6px;
  }
  .page-placeholder { display: flex; align-items: center; justify-content: center; color: var(--secondary); font-size: 12px; }
  .page-meta { font-size: 11px; color: var(--secondary); margin-top: 5px; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
  .page-delete {
    position: absolute; top: 4px; right: 4px; width: 22px; height: 22px;
    border: none; border-radius: 6px; background: rgba(255,255,255,0.94);
    color: #ff3b30; font-size: 15px; line-height: 1; cursor: pointer;
    box-shadow: 0 1px 3px rgba(0,0,0,0.18);
  }
  .page-delete:hover { background: #fff0f0; }
  .page-empty {
    grid-column: 1 / -1; text-align: center;
    color: var(--secondary); font-size: 14px; padding: 24px 0;
  }
  .file-item .file-info { font-size: 12px; color: var(--secondary); white-space: nowrap; }

  /* Lightbox */
  .lightbox {
    display: none; position: fixed; inset: 0; z-index: 50;
    background: rgba(0,0,0,0.55); align-items: center; justify-content: center;
    padding: 30px;
  }
  .lightbox.open { display: flex; }
  .lightbox img {
    max-width: 92vw; max-height: 88vh; background: #fff;
    border-radius: 8px; box-shadow: 0 20px 60px rgba(0,0,0,0.35);
  }
  .lightbox-close {
    position: absolute; top: 18px; right: 22px; border: none;
    background: rgba(255,255,255,0.92); border-radius: 8px;
    padding: 8px 14px; font-size: 14px; cursor: pointer;
  }

  /* Footer */
  .footer {
    text-align: center; color: var(--secondary);
    font-size: 12px; margin-top: 40px;
  }

  input[type="file"] { display: none; }
</style>
</head>
<body>
<div class="container">
  <h1>换墨</h1>
  <p class="subtitle">Markdown / TXT / DOCX → PDF · PDF 合并 · 中文字体自动检测</p>

  <div class="mode-switch" id="modeSwitch">
    <button type="button" class="mode-btn active" data-mode="convert">文档转 PDF</button>
    <button type="button" class="mode-btn" data-mode="merge">合并 PDF</button>
  </div>

  <!-- Drop zone -->
  <div class="drop-zone" id="dropZone">
    <div class="icon"><svg width="36" height="36" viewBox="0 0 24 24" fill="none" stroke="#86868b" stroke-width="1.5" stroke-linecap="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="12" y1="18" x2="12" y2="12"/><polyline points="9 15 12 18 15 15"/></svg></div>
    <p>拖拽 <strong id="dropHint">.md / .txt / .docx 文件</strong> 到此处<br>或 <span class="browse">点击浏览</span></p>
  </div>
  <input type="file" id="fileInput" accept=".md,.markdown,.txt,.text,.docx,.pdf" multiple>

  <!-- File list -->
  <div class="file-list" id="fileList">
    <div class="file-list-header">已选文件 (<span id="fileCount">0</span>)</div>
    <div id="fileItems"></div>
  </div>

  <!-- Page editor (merge mode) -->
  <div class="page-editor" id="pageEditor">
    <div class="page-editor-header">合并页面（<span id="pageCount">0</span> 页）· 拖拽排序 · 点击 × 删除</div>
    <div class="page-list" id="pageList"></div>
  </div>

  <!-- Output dir -->
  <div class="output-row">
    <label for="outputDir">输出到</label>
    <input type="text" id="outputDir" value="{{ output_dir }}">
    <button class="btn-outline" id="pickOutput">选择…</button>
  </div>

  <!-- Convert -->
  <button class="btn-convert" id="btnConvert" disabled>
    <span class="spinner">···</span>
    <span class="btn-text" id="btnText">开始转换</span>
  </button>

  <!-- Results -->
  <div class="results" id="results">
    <div class="results-header" id="resultsHeader">转换结果</div>
    <div id="resultItems"></div>
  </div>

  <p class="footer">换墨 · macOS / Linux / Windows</p>
</div>

<div class="lightbox" id="lightbox">
  <button class="lightbox-close" id="lightboxClose">关闭</button>
  <img id="lightboxImg" alt="页面预览">
</div>

<script>
const dropZone = document.getElementById('dropZone');
const fileInput = document.getElementById('fileInput');
const fileList = document.getElementById('fileList');
const fileCount = document.getElementById('fileCount');
const fileItems = document.getElementById('fileItems');
const btnConvert = document.getElementById('btnConvert');
const outputDir = document.getElementById('outputDir');
const results = document.getElementById('results');
const resultItems = document.getElementById('resultItems');
const modeSwitch = document.getElementById('modeSwitch');
const dropHint = document.getElementById('dropHint');
const btnText = document.getElementById('btnText');
const resultsHeader = document.getElementById('resultsHeader');
const pageEditor = document.getElementById('pageEditor');
const pageList = document.getElementById('pageList');
const pageCountEl = document.getElementById('pageCount');
const lightbox = document.getElementById('lightbox');
const lightboxImg = document.getElementById('lightboxImg');
const lightboxClose = document.getElementById('lightboxClose');

let files = new Map(); // name -> content (base64)
let mode = 'convert';
let pdfMeta = new Map(); // name -> {pages, token, preview}
let pageOrder = [];      // {id, file, page}
let dragId = null;
let dropTargetId = null;
let dropAfter = false;

function setMode(nextMode) {
  mode = nextMode;
  modeSwitch.querySelectorAll('.mode-btn').forEach(btn => {
    btn.classList.toggle('active', btn.dataset.mode === mode);
  });

  if (mode === 'merge') {
    dropHint.textContent = '.pdf 文件（可多选）';
    fileInput.accept = '.pdf';
    btnText.textContent = '开始合并';
    resultsHeader.textContent = '合并结果';
    pageEditor.classList.add('visible');
  } else {
    dropHint.textContent = '.md / .txt / .docx 文件';
    fileInput.accept = '.md,.markdown,.txt,.text,.docx';
    btnText.textContent = '开始转换';
    resultsHeader.textContent = '转换结果';
    pageEditor.classList.remove('visible');
  }

  files.clear();
  pdfMeta.clear();
  pageOrder = [];
  renderFileList();
  renderMergeList();
  resultItems.innerHTML = '';
  results.classList.remove('has-results');
}

modeSwitch.querySelectorAll('.mode-btn').forEach(btn => {
  btn.addEventListener('click', () => setMode(btn.dataset.mode));
});

// ── File handling ─────────────────────────────────────────
dropZone.addEventListener('click', () => fileInput.click());

dropZone.addEventListener('dragover', e => { e.preventDefault(); dropZone.classList.add('drag-over'); });
dropZone.addEventListener('dragleave', () => dropZone.classList.remove('drag-over'));
dropZone.addEventListener('drop', e => {
  e.preventDefault();
  dropZone.classList.remove('drag-over');
  addFiles(e.dataTransfer.files);
});

fileInput.addEventListener('change', () => {
  addFiles(fileInput.files);
  fileInput.value = '';
});

async function addFiles(fileList) {
  if (mode === 'merge') {
    await addPdfFiles(fileList);
    return;
  }

  for (const f of fileList) {
    if (!isAccepted(f.name)) continue;
    const reader = new FileReader();
    reader.onload = () => {
      files.set(f.name, reader.result.split(',')[1]); // base64 content
      renderFileList();
    };
    reader.readAsDataURL(f);
  }
}

async function addPdfFiles(fileList) {
  for (const f of fileList) {
    if (!f.name.toLowerCase().endsWith('.pdf')) continue;
    if (files.has(f.name)) {
      alert(`已存在同名文件：${f.name}`);
      continue;
    }

    const reader = new FileReader();
    const b64 = await new Promise(resolve => {
      reader.onload = () => resolve(reader.result.split(',')[1]);
      reader.readAsDataURL(f);
    });

    files.set(f.name, b64);
    renderFileList();

    try {
      const resp = await fetch('/analyze-pdf', {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ name: f.name, content: b64 }),
      });
      const data = await resp.json();
      if (!resp.ok) throw new Error(data.error || '读取失败');

      pdfMeta.set(f.name, data);
      for (let i = 0; i < data.pages; i++) {
        pageOrder.push({ id: `${f.name}::${i}`, file: f.name, page: i });
      }
    } catch (err) {
      alert(`${f.name} 读取失败：${err.message}`);
      files.delete(f.name);
      pdfMeta.delete(f.name);
      pageOrder = pageOrder.filter(p => p.file !== f.name);
    }

    renderFileList();
    renderMergeList();
  }
}

function isAccepted(name) {
  const lower = name.toLowerCase();
  if (mode === 'merge') return lower.endsWith('.pdf');
  return lower.endsWith('.md') || lower.endsWith('.markdown') || lower.endsWith('.txt') || lower.endsWith('.text') || lower.endsWith('.docx');
}

function renderFileList() {
  fileItems.innerHTML = '';
  let i = 0;
  for (const [name] of files) {
    const div = document.createElement('div');
    div.className = 'file-item';
    let extra = '';
    if (mode === 'merge') {
      const meta = pdfMeta.get(name);
      const kept = pageOrder.filter(p => p.file === name).length;
      extra = meta ? `${meta.pages} 页 · 保留 ${kept} 页` : '读取中…';
    }
    div.innerHTML = `
      <div class="name"><span>${esc(name)}</span><span class="file-info">${extra}</span></div>
      <button class="remove" data-name="${esc(name)}">移除</button>
    `;
    fileItems.appendChild(div);
    i++;
  }
  fileCount.textContent = files.size;
  const ready = mode === 'merge' ? (files.size >= 2 && pageOrder.length > 0) : files.size > 0;
  fileList.classList.toggle('has-files', files.size > 0);
  btnConvert.disabled = !ready;

  // Bind remove buttons
  fileItems.querySelectorAll('.remove').forEach(btn => {
    btn.addEventListener('click', () => {
      files.delete(btn.dataset.name);
      pdfMeta.delete(btn.dataset.name);
      pageOrder = pageOrder.filter(p => p.file !== btn.dataset.name);
      renderFileList();
      renderMergeList();
    });
  });
}

function renderMergeList() {
  pageCountEl.textContent = pageOrder.length;
  pageList.innerHTML = '';
  if (pageOrder.length === 0) {
    pageList.innerHTML = '<div class="page-empty">暂无页面，请先添加 PDF 文件</div>';
    return;
  }

  for (const p of pageOrder) {
    const meta = pdfMeta.get(p.file);
    const card = document.createElement('div');
    card.className = 'page-card';
    card.draggable = true;
    card.dataset.id = p.id;

    let preview;
    if (meta && meta.preview) {
      preview = `<img src="/page-preview/${encodeURIComponent(meta.token)}/${p.page}" alt="${esc(p.file)} 第 ${p.page + 1} 页" loading="lazy">`;
    } else {
      const reason = meta && meta.preview_error ? ` title="${esc(meta.preview_error)}"` : '';
      preview = `<div class="page-placeholder"${reason}>第 ${p.page + 1} 页</div>`;
    }

    card.innerHTML = `
      ${preview}
      <button class="page-delete" title="删除此页">×</button>
      <div class="page-meta">${esc(p.file)} · ${p.page + 1} 页</div>
    `;
    pageList.appendChild(card);
  }
}

// ── Page reorder & preview (merge mode) ──────────────────
pageList.addEventListener('dragstart', e => {
  const card = e.target.closest('.page-card');
  if (!card) return;
  dragId = card.dataset.id;
  dropTargetId = null;
  dropAfter = false;
  card.classList.add('dragging');
  e.dataTransfer.effectAllowed = 'move';
});

pageList.addEventListener('dragover', e => {
  e.preventDefault();
  const card = e.target.closest('.page-card');
  if (!card || card.dataset.id === dragId) return;

  pageList.querySelectorAll('.page-card').forEach(c => c.classList.remove('drop-target'));
  card.classList.add('drop-target');
  const rect = card.getBoundingClientRect();
  dropAfter = (e.clientY - rect.top) > rect.height / 2;
  dropTargetId = card.dataset.id;
});

pageList.addEventListener('drop', e => {
  e.preventDefault();
  pageList.querySelectorAll('.page-card').forEach(c => c.classList.remove('drop-target'));

  const dragged = pageOrder.find(p => p.id === dragId);
  const target = pageOrder.find(p => p.id === dropTargetId);
  if (dragged && target && dragId !== dropTargetId) {
    pageOrder = pageOrder.filter(p => p.id !== dragId);
    const idx = pageOrder.indexOf(target);
    pageOrder.splice(dropAfter ? idx + 1 : idx, 0, dragged);
  }

  dragId = null;
  dropTargetId = null;
  dropAfter = false;
  renderMergeList();
  renderFileList();
});

pageList.addEventListener('dragend', () => {
  dragId = null;
  dropTargetId = null;
  dropAfter = false;
  pageList.querySelectorAll('.page-card').forEach(c => c.classList.remove('dragging', 'drop-target'));
  renderMergeList();
  renderFileList();
});

pageList.addEventListener('click', e => {
  const del = e.target.closest('.page-delete');
  if (del) {
    const id = del.closest('.page-card').dataset.id;
    pageOrder = pageOrder.filter(p => p.id !== id);
    renderMergeList();
    renderFileList();
    return;
  }

  const card = e.target.closest('.page-card');
  if (!card) return;
  const p = pageOrder.find(x => x.id === card.dataset.id);
  const meta = p && pdfMeta.get(p.file);
  if (p && meta && meta.preview) {
    lightboxImg.src = `/page-preview/${encodeURIComponent(meta.token)}/${p.page}`;
    lightbox.classList.add('open');
  }
});

lightboxClose.addEventListener('click', () => lightbox.classList.remove('open'));
lightbox.addEventListener('click', e => {
  if (e.target === lightbox) lightbox.classList.remove('open');
});

// ── Convert ────────────────────────────────────────────────
btnConvert.addEventListener('click', async () => {
  if (files.size === 0) return;
  if (mode === 'merge') {
    mergeFiles();
    return;
  }

  btnConvert.classList.add('loading');
  btnConvert.disabled = true;
  results.classList.remove('has-results');
  resultItems.innerHTML = '';

  const dir = outputDir.value.trim() || '{{ output_dir }}';

  for (const [name, b64] of files) {
    const itemDiv = document.createElement('div');
    itemDiv.className = 'result-item';
    itemDiv.innerHTML = `<span class="status">·</span><div class="info"><div class="filename">${esc(name)}</div><div class="size">转换中…</div></div>`;
    resultItems.appendChild(itemDiv);

    try {
      const resp = await fetch('/convert', {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ name, content: b64, output_dir: dir }),
      });

      if (!resp.ok) {
        const err = await resp.json();
        throw new Error(err.error || '未知错误');
      }

      const data = await resp.json();
      itemDiv.innerHTML = `
        <span class="status">&#10003;</span>
        <div class="info"><div class="filename">${esc(data.filename)}</div><div class="size">${data.size_kb} KB</div></div>
        <a class="open-btn" href="/download/${encodeURIComponent(data.token)}" target="_blank">打开</a>
      `;
    } catch (err) {
      itemDiv.classList.add('error');
      itemDiv.innerHTML = `
        <span class="status">&#10007;</span>
        <div class="info"><div class="filename">${esc(name)}</div><div class="size">${esc(err.message)}</div></div>
      `;
    }
  }

  results.classList.add('has-results');
  btnConvert.classList.remove('loading');
  btnConvert.disabled = false;
});

async function mergeFiles() {
  if (files.size < 2) {
    alert('请至少选择两个 PDF 文件');
    btnConvert.disabled = !(files.size >= 2 && pageOrder.length > 0);
    return;
  }
  if (pageOrder.length === 0) {
    alert('请至少保留一个页面');
    btnConvert.disabled = true;
    return;
  }

  btnConvert.classList.add('loading');
  btnConvert.disabled = true;
  results.classList.remove('has-results');
  resultItems.innerHTML = '';

  const dir = outputDir.value.trim() || '{{ output_dir }}';
  const payload = {
    files: [...files.entries()].map(([name, content]) => ({ name, content })),
    order: pageOrder.map(p => ({ file: p.file, page: p.page })),
    output_dir: dir,
  };

  const itemDiv = document.createElement('div');
  itemDiv.className = 'result-item';
  itemDiv.innerHTML = `<span class="status">·</span><div class="info"><div class="filename">合并中…</div><div class="size">按当前顺序合并 ${pageOrder.length} 页</div></div>`;
  resultItems.appendChild(itemDiv);

  try {
    const resp = await fetch('/merge-pdf', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify(payload),
    });
    const data = await resp.json();
    if (!resp.ok) throw new Error(data.error || '合并失败');

    itemDiv.innerHTML = `
      <span class="status">&#10003;</span>
      <div class="info"><div class="filename">${esc(data.filename)}</div><div class="size">${data.size_kb} KB</div></div>
      <a class="open-btn" href="/download/${encodeURIComponent(data.token)}" target="_blank">打开</a>
    `;
  } catch (err) {
    itemDiv.classList.add('error');
    itemDiv.innerHTML = `
      <span class="status">&#10007;</span>
      <div class="info"><div class="filename">合并失败</div><div class="size">${esc(err.message)}</div></div>
    `;
  }

  results.classList.add('has-results');
  btnConvert.classList.remove('loading');
  btnConvert.disabled = false;
}

// ── Output picker ─────────────────────────────────────────
document.getElementById('pickOutput').addEventListener('click', async () => {
  try {
    const resp = await fetch('/pick-folder', { method: 'POST' });
    const data = await resp.json();
    if (data.path) outputDir.value = data.path;
  } catch (err) {
    alert('选择目录出错，请手动输入路径');
  }
});

outputDir.addEventListener('change', () => updateOutputDir());
function updateOutputDir() {
  fetch('/set-output-dir', {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ path: outputDir.value.trim() }),
  });
}

function esc(s) {
  return s.replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;').replace(/"/g, '&quot;');
}
</script>
</body>
</html>"""

# ── Routes ────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE, output_dir=OUTPUT_DIR)


@app.route("/convert", methods=["POST"])
def convert():
    data = request.get_json()
    name = data.get("name", "untitled.md")
    b64_content = data.get("content", "")
    output_dir = data.get("output_dir", OUTPUT_DIR)

    ext = Path(name).suffix.lower()

    # Decode base64 — text for md/txt, binary for docx
    try:
        raw_bytes = base64.b64decode(b64_content)
        if ext == ".docx":
            text = None  # binary, handled separately
        else:
            text = raw_bytes.decode("utf-8")
    except Exception as e:
        return jsonify({"error": f"文件解码失败: {e}"}), 400

    # Font
    font = detect_cjk_font()
    if not font:
        return jsonify({"error": "未检测到中文字体，请安装字体后重试"}), 500

    # Output dir
    out_path = Path(output_dir)
    try:
        out_path.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        return jsonify({"error": f"无法创建输出目录: {e}"}), 400

    # Determine file type and convert
    title = Path(name).stem.replace("_", " ")

    try:
        if ext == ".txt":
            pdf_bytes = convert_txt_to_pdf_bytes(text, title, font)
        elif ext == ".docx":
            pdf_bytes = convert_docx_to_pdf_bytes(raw_bytes, title, font)
        else:
            pdf_bytes = convert_md_to_pdf_bytes(text, title, font)
    except Exception as e:
        return jsonify({"error": f"转换失败: {e}"}), 500

    # Save to output dir
    pdf_name = Path(name).with_suffix(".pdf").name
    pdf_path = out_path / pdf_name
    try:
        pdf_path.write_bytes(pdf_bytes)
    except OSError as e:
        return jsonify({"error": f"写入文件失败: {e}"}), 500

    # Store in a simple in-memory registry for /download
    token = base64.urlsafe_b64encode(str(pdf_path).encode()).decode()
    _downloads[token] = str(pdf_path)

    return jsonify({
        "filename": pdf_name,
        "size_kb": round(len(pdf_bytes) / 1024),
        "token": token,
    })


_downloads: dict[str, str] = {}


_PREVIEW_DIR = Path(tempfile.gettempdir()) / "huanmo_previews"
_PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
_preview_cache: dict[str, dict] = {}


@app.route("/analyze-pdf", methods=["POST"])
def analyze_pdf():
    data = request.get_json()
    name = data.get("name", "file.pdf")
    b64_content = data.get("content", "")

    try:
        raw = base64.b64decode(b64_content)
        reader = PdfReader(BytesIO(raw))
        if getattr(reader, "is_encrypted", False):
            return jsonify({"error": f"{name} 已加密，无法预览"}), 400
        page_count = len(reader.pages)
        if page_count == 0:
            return jsonify({"error": f"{name} 没有可预览的页面"}), 400
    except Exception as e:
        return jsonify({"error": f"读取失败: {e}"}), 400

    token = base64.urlsafe_b64encode(os.urandom(9)).decode("ascii").rstrip("=")
    preview_dir = _PREVIEW_DIR / token
    preview_ok = False
    preview_error = ""
    try:
        preview_dir.mkdir(parents=True, exist_ok=True)
        for i in range(page_count):
            png = render_page_preview(raw, i)
            if png is None:
                preview_error = "预览渲染库不可用，请确认已安装 pypdfium2"
                break
            (preview_dir / f"{i}.png").write_bytes(png)
        preview_ok = (preview_dir / "0.png").is_file()
    except Exception as e:
        preview_ok = False
        preview_error = f"{type(e).__name__}: {e}"
        log.warning("PDF 预览失败 %s: %s", name, preview_error)

    _preview_cache[token] = {"name": name, "dir": str(preview_dir)}
    return jsonify({
        "name": name,
        "pages": page_count,
        "token": token,
        "preview": preview_ok,
        "preview_error": preview_error,
    })


@app.route("/page-preview/<token>/<int:page>")
def page_preview(token, page):
    info = _preview_cache.get(token)
    if not info:
        return "预览不存在或已过期", 404
    path = Path(info["dir"]) / f"{page}.png"
    if not path.is_file():
        return "预览不存在", 404
    return send_file(path, mimetype="image/png")


@app.route("/merge-pdf", methods=["POST"])
def merge_pdf():
    data = request.get_json()
    items = data.get("files", [])
    order = data.get("order")
    output_dir = data.get("output_dir", OUTPUT_DIR)

    if len(items) < 2:
        return jsonify({"error": "请至少选择两个 PDF 文件"}), 400

    try:
        pdf_items = [
            (item.get("name", "file.pdf"), base64.b64decode(item.get("content", "")))
            for item in items
        ]
        if order is not None:
            if not order:
                return jsonify({"error": "请至少保留一个页面"}), 400
            ordered: list[tuple[str, int]] = []
            readers = {name: PdfReader(BytesIO(raw)) for name, raw in pdf_items}
            for entry in order:
                file_name = entry.get("file")
                page_index = entry.get("page")
                if file_name not in readers:
                    raise ValueError(f"未知文件: {file_name}")
                if not isinstance(page_index, int) or not 0 <= page_index < len(readers[file_name].pages):
                    raise ValueError(f"{file_name} 页面索引无效")
                ordered.append((file_name, page_index))
            pdf_bytes = merge_selected_pages(pdf_items, ordered)
        else:
            pdf_bytes = merge_pdf_bytes(pdf_items)
    except Exception as e:
        return jsonify({"error": f"合并失败: {e}"}), 500

    out_path = Path(output_dir)
    try:
        out_path.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        return jsonify({"error": f"无法创建输出目录: {e}"}), 400

    first_stem = Path(items[0].get("name", "merged")).stem
    pdf_name = f"{first_stem}_合并.pdf"
    pdf_path = out_path / pdf_name
    try:
        pdf_path.write_bytes(pdf_bytes)
    except OSError as e:
        return jsonify({"error": f"写入文件失败: {e}"}), 500

    token = base64.urlsafe_b64encode(str(pdf_path).encode()).decode()
    _downloads[token] = str(pdf_path)

    return jsonify({
        "filename": pdf_name,
        "size_kb": round(len(pdf_bytes) / 1024),
        "token": token,
        "files": len(items),
    })


@app.route("/download/<token>")
def download(token):
    path = _downloads.get(token)
    if not path or not os.path.isfile(path):
        return "文件不存在或已过期", 404
    return send_file(path, as_attachment=True)


@app.route("/set-output-dir", methods=["POST"])
def set_output_dir():
    global OUTPUT_DIR
    data = request.get_json()
    p = data.get("path", "").strip()
    if p and os.path.isdir(os.path.dirname(p) if not os.path.isdir(p) else p):
        OUTPUT_DIR = p
    return jsonify({"ok": True})


@app.route("/pick-folder", methods=["POST"])
def pick_folder():
    """Use AppleScript on macOS to open a folder picker dialog."""
    global OUTPUT_DIR
    if sys.platform != "darwin":
        return jsonify({"path": OUTPUT_DIR})
    import subprocess
    script = '''
        tell application "System Events"
            activate
            set f to choose folder with prompt "选择 PDF 输出目录:" default location path to downloads folder
            POSIX path of f
        end tell
    '''
    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=30,
        )
        path = result.stdout.strip()
        if path:
            OUTPUT_DIR = path
            return jsonify({"path": path})
    except Exception:
        pass
    return jsonify({"path": OUTPUT_DIR})


# ── Main ──────────────────────────────────────────────────────

def main():
    logging.basicConfig(level=logging.INFO, format="%(message)s")

    port = 5199
    url = f"http://127.0.0.1:{port}"

    print("╔══════════════════════════════════════╗")
    print("║      换墨 — 文档转换 + PDF 合并    ║")
    print("╠══════════════════════════════════════╣")
    print(f"║  浏览器地址: {url}          ║")
    print("║  按 Ctrl+C 停止服务                  ║")
    print("╚══════════════════════════════════════╝")

    # Auto-open browser after a short delay
    threading.Timer(1.0, lambda: webbrowser.open(url)).start()

    app.run(host="127.0.0.1", port=port, debug=False)


if __name__ == "__main__":
    main()
